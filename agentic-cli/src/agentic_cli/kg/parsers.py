"""Parsers for different data sources."""

import csv
import json
import re
from pathlib import Path
from typing import Any, Dict, List, Optional, Set


def parse_pdf(file_path: str) -> List[Dict[str, Any]]:
    """
    Parse PDF file and extract text content.
    
    Args:
        file_path: Path to PDF file
    
    Returns:
        List of document dictionaries
    """
    try:
        import PyPDF2
    except ImportError:
        raise ImportError(
            "PyPDF2 not installed. Install with: pip install PyPDF2"
        )
    
    path = Path(file_path)
    if not path.exists():
        raise FileNotFoundError(f"File not found: {file_path}")
    
    documents = []
    
    try:
        with open(path, "rb") as f:
            reader = PyPDF2.PdfReader(f)
            
            for i, page in enumerate(reader.pages):
                text = page.extract_text()
                
                if text.strip():
                    documents.append({
                        "title": f"{path.stem} - Page {i + 1}",
                        "content": text,
                        "metadata": {
                            "page": i + 1,
                            "total_pages": len(reader.pages),
                            "format": "pdf",
                        },
                    })
    except Exception as e:
        # Log error but return empty list instead of raising
        import logging
        logger = logging.getLogger(__name__)
        logger.warning(f"Failed to parse PDF {path.name}: {e}")
        return []
    
    return documents


def parse_text(file_path: str, metadata: Optional[Dict[str, Any]] = None) -> List[Dict[str, Any]]:
    """
    Parse text file.
    
    Args:
        file_path: Path to text file
        metadata: Optional metadata to merge with document metadata
    
    Returns:
        List of document dictionaries
    """
    path = Path(file_path)
    if not path.exists():
        raise FileNotFoundError(f"File not found: {file_path}")
    
    with open(path, "r", encoding="utf-8") as f:
        content = f.read()
    
    doc_metadata = {
        "format": "text",
        "size": len(content),
    }
    
    # Merge provided metadata
    if metadata:
        doc_metadata.update(metadata)
    
    # Use document_title from metadata if available, otherwise use file stem
    title = doc_metadata.get("document_title", path.stem)
    
    return [{
        "title": title,
        "content": content,
        "metadata": doc_metadata,
    }]


def parse_csv(file_path: str) -> List[Dict[str, Any]]:
    """
    Parse CSV file.
    
    Args:
        file_path: Path to CSV file
    
    Returns:
        List of document dictionaries (one per row)
    """
    path = Path(file_path)
    if not path.exists():
        raise FileNotFoundError(f"File not found: {file_path}")
    
    documents = []
    
    with open(path, "r", encoding="utf-8") as f:
        reader = csv.DictReader(f)
        
        for i, row in enumerate(reader):
            # Convert row to text content
            content = "\n".join([f"{k}: {v}" for k, v in row.items()])
            
            documents.append({
                "title": f"{path.stem} - Row {i + 1}",
                "content": content,
                "metadata": {
                    "row": i + 1,
                    "format": "csv",
                    "fields": list(row.keys()),
                    "data": row,
                },
            })
    
    return documents


def parse_json(file_path: str) -> List[Dict[str, Any]]:
    """
    Parse JSON file.
    
    Args:
        file_path: Path to JSON file
    
    Returns:
        List of document dictionaries
    """
    path = Path(file_path)
    if not path.exists():
        raise FileNotFoundError(f"File not found: {file_path}")
    
    with open(path, "r", encoding="utf-8") as f:
        data = json.load(f)
    
    documents = []
    
    # Handle different JSON structures
    if isinstance(data, list):
        # List of objects
        for i, item in enumerate(data):
            content = json.dumps(item, indent=2)
            
            documents.append({
                "title": f"{path.stem} - Item {i + 1}",
                "content": content,
                "metadata": {
                    "index": i,
                    "format": "json",
                    "data": item,
                },
            })
    elif isinstance(data, dict):
        # Single object or nested structure
        content = json.dumps(data, indent=2)
        
        documents.append({
            "title": path.stem,
            "content": content,
            "metadata": {
                "format": "json",
                "data": data,
            },
        })
    else:
        # Primitive value
        documents.append({
            "title": path.stem,
            "content": str(data),
            "metadata": {
                "format": "json",
                "data": data,
            },
        })
    
    return documents


def _html_to_text(html: str) -> str:
    """Convert Confluence storage-format HTML to clean text/markdown."""
    import re

    text = html
    # Block-level elements → newlines
    text = re.sub(r'<br\s*/?>', '\n', text)
    text = re.sub(r'</?(p|div|tr|li|blockquote)[^>]*>', '\n', text)
    text = re.sub(r'</(h[1-6])[^>]*>', '\n\n', text)
    # Headers
    for level in range(1, 7):
        text = re.sub(rf'<h{level}[^>]*>', '\n' + '#' * level + ' ', text)
    # Bold / italic
    text = re.sub(r'<(strong|b)[^>]*>(.*?)</\1>', r'**\2**', text)
    text = re.sub(r'<(em|i)[^>]*>(.*?)</\1>', r'*\2*', text)
    # Code blocks
    text = re.sub(r'<ac:structured-macro[^>]*ac:name="code"[^>]*>.*?<ac:plain-text-body>\s*<!\[CDATA\[(.*?)\]\]>\s*</ac:plain-text-body>\s*</ac:structured-macro>', r'\n```\n\1\n```\n', text, flags=re.DOTALL)
    # Tables — keep as plain text rows
    text = re.sub(r'<t[hd][^>]*>', ' | ', text)
    text = re.sub(r'</t[hd]>', '', text)
    # Links
    text = re.sub(r'<a[^>]*href="([^"]*)"[^>]*>(.*?)</a>', r'[\2](\1)', text)
    # Strip remaining tags
    text = re.sub(r'<[^>]+>', '', text)
    # HTML entities
    text = text.replace('&amp;', '&').replace('&lt;', '<').replace('&gt;', '>').replace('&nbsp;', ' ').replace('&quot;', '"')
    # Collapse blank lines
    text = re.sub(r'\n{3,}', '\n\n', text)
    return text.strip()


def _extract_referenced_attachments(body_html: str) -> Set[str]:
    """
    Extract attachment filenames referenced in Confluence page content.
    
    Confluence storage format uses: <ac:link><ri:attachment ri:filename="..." /></ac:link>
    
    Args:
        body_html: Page content in Confluence storage format
    
    Returns:
        Set of attachment filenames referenced in the content
    """
    referenced = set()
    # Match Confluence attachment references in storage format
    pattern = r'<ri:attachment\s+ri:filename="([^"]+)"'
    matches = re.findall(pattern, body_html)
    referenced.update(matches)
    return referenced


def _get_confluence_client(config=None):
    """Create and return a Confluence client + base_url tuple."""
    try:
        from atlassian import Confluence
    except ImportError:
        raise ImportError(
            "atlassian-python-api not installed. Install with: pip install atlassian-python-api"
        )

    if config is None:
        from agentic_cli.kg.config import KGConfig
        config = KGConfig.load()

    if not config.is_confluence_configured():
        raise ValueError(
            "Confluence is not configured. Please run:\n"
            "  agent-cli kg init --confluence-url <url> --confluence-username <username> --confluence-token <token>"
        )

    from urllib.parse import urlparse
    base_url = config.confluence_url
    if "/spaces/" in base_url or "/pages/" in base_url:
        parsed = urlparse(base_url)
        base_url = f"{parsed.scheme}://{parsed.netloc}"

    confluence = Confluence(
        url=base_url,
        username=config.confluence_username,
        password=config.confluence_api_token,
        cloud=config.confluence_cloud,
    )
    return confluence, base_url, config


def _extract_page_id_from_url(url: str) -> str:
    """Extract the numeric page ID from a Confluence URL."""
    if "/pages/" in url:
        return url.split("/pages/")[1].split("/")[0]
    raise ValueError(f"Cannot extract page ID from URL: {url}")


def _fetch_page_tree(confluence, page_id: str, max_depth: int = 3, _depth: int = 0) -> list:
    """Recursively fetch a page and all its children."""
    pages = []

    page = confluence.get_page_by_id(page_id, expand='body.storage,version,space,children.page')
    pages.append(page)

    if _depth < max_depth:
        children = page.get("children", {}).get("page", {}).get("results", [])
        for child in children:
            try:
                pages.extend(_fetch_page_tree(confluence, child["id"], max_depth, _depth + 1))
            except Exception:
                continue  # skip inaccessible child pages

    return pages


def parse_confluence(url: str, config=None) -> List[Dict[str, Any]]:
    """
    Parse Confluence page or space.
    
    Args:
        url: Confluence page or space URL
        config: KGConfig instance with Confluence credentials
    
    Returns:
        List of document dictionaries
    """
    confluence, base_url, config = _get_confluence_client(config)
    
    documents = []
    
    try:
        # Parse URL to extract page ID or space key
        if "/pages/" in url:
            # Extract page ID from URL
            page_id = url.split("/pages/")[1].split("/")[0]
            
            # Get page content
            page = confluence.get_page_by_id(page_id, expand='body.storage,version')
            
            documents.append({
                "content": _html_to_text(page['body']['storage']['value']),
                "metadata": {
                    "source": url,
                    "title": page['title'],
                    "space": page['space']['key'],
                    "page_id": page_id,
                    "version": page['version']['number'],
                    "type": "confluence_page"
                }
            })
            
        elif "/spaces/" in url:
            # Extract space key from URL
            space_key = url.split("/spaces/")[1].split("/")[0]
            
            # Get all pages in the space
            pages = confluence.get_all_pages_from_space(space_key, expand='body.storage,version')
            
            for page in pages:
                documents.append({
                    "content": _html_to_text(page['body']['storage']['value']),
                    "metadata": {
                        "source": f"{config.confluence_url}/pages/{page['id']}",
                        "title": page['title'],
                        "space": space_key,
                        "page_id": page['id'],
                        "version": page['version']['number'],
                        "type": "confluence_page"
                    }
                })
        else:
            raise ValueError(
                f"Unsupported Confluence URL format: {url}\n"
                "Supported formats:\n"
                "  - Page: https://confluence.company.com/pages/123456/Page+Title\n"
                "  - Space: https://confluence.company.com/spaces/SPACEKEY"
            )
    
    except Exception as e:
        raise RuntimeError(f"Failed to fetch Confluence content: {str(e)}")
    
    return documents


def parse_confluence_tree(
    url: str,
    config=None,
    include_children: bool = True,
    max_depth: int = 3,
    latest_only: bool = False,
    use_mcp: bool = False,
    include_attachments: bool = True,
) -> List[Dict[str, Any]]:
    """
    Parse a Confluence page and its entire child-page tree.

    Designed for release-style hierarchical pages like:
        Release 29/
        ├── Architecture Changes
        ├── API Reference
        └── Deployment Notes

    Args:
        url: Confluence page URL (must contain /pages/<id>/)
        config: KGConfig instance with Confluence credentials
        include_children: Recursively fetch child pages (default True)
        max_depth: Maximum depth for child page crawling (default 3)
        latest_only: If True, only return the latest version of each page
        use_mcp: If True, use MCP Confluence server instead of direct API (bypasses KG config check)
        include_attachments: If True, fetch and parse page attachments (default True)

    Returns:
        List of document dictionaries (parent + children)
    """
    # Check MCP mode first (bypasses all config validation)
    if use_mcp:
        # Use MCP-based Confluence client (bypasses KG config check)
        return _parse_confluence_tree_mcp(url, include_children=include_children, max_depth=max_depth, include_attachments=include_attachments, parent_page_id=None)
    
    confluence, base_url, cfg = _get_confluence_client(config)

    page_id = _extract_page_id_from_url(url)

    if include_children:
        raw_pages = _fetch_page_tree(confluence, page_id, max_depth=max_depth)
    else:
        raw_pages = [confluence.get_page_by_id(page_id, expand='body.storage,version,space')]

    documents = []
    seen_ids = set()

    for page in raw_pages:
        pid = str(page.get("id", ""))
        if pid in seen_ids:
            continue
        seen_ids.add(pid)

        html_body = page.get("body", {}).get("storage", {}).get("value", "")
        title = page.get("title", f"Page {pid}")
        version = page.get("version", {}).get("number", 0)
        space_key = page.get("space", {}).get("key", "")

        # Convert storage HTML to clean text/markdown
        clean_text = _html_to_text(html_body)
        if not clean_text.strip():
            continue

        # Build a richer document with header for KG ingestion
        doc_text = f"# {title}\n\n"
        doc_text += f"Source: Confluence page {pid} (space: {space_key}, version: {version})\n\n"
        doc_text += clean_text

        documents.append({
            "content": doc_text,
            "metadata": {
                "source": f"{base_url}/pages/{pid}",
                "title": title,
                "space": space_key,
                "page_id": pid,
                "version": version,
                "type": "confluence_page",
                "parent_page_id": page_id if pid != page_id else None,
            }
        })

    return documents


def _parse_confluence_tree_mcp(
    url: str,
    include_children: bool = True,
    max_depth: int = 3,
    include_attachments: bool = True,
    current_depth: int = 0,
    parent_page_id: Optional[str] = None,
) -> List[Dict[str, Any]]:
    """
    Parse Confluence page tree using MCP server (bypasses KG config check).
    
    Args:
        url: Confluence page URL
        include_children: Whether to fetch child pages
        max_depth: Maximum depth for child page crawling
        include_attachments: Whether to fetch and parse attachments
        current_depth: Current recursion depth
        parent_page_id: Parent page ID (for linking child pages to release pages)
        
    Returns:
        List of document dictionaries
    """
    from agentic_cli.mcp_tool_client import confluence_get_page, confluence_list_attachments, confluence_download_attachment, confluence_get_page_children, MCPToolError
    from urllib.parse import urlparse
    import tempfile
    import os
    import base64
    import logging
    
    logger = logging.getLogger(__name__)
    
    # Extract page ID from URL
    if "/pages/" not in url:
        raise ValueError(f"Invalid Confluence page URL: {url}")
    
    parsed = urlparse(url)
    path_parts = parsed.path.split("/")
    try:
        page_idx = path_parts.index("pages")
        page_id = path_parts[page_idx + 1]
    except (ValueError, IndexError):
        raise ValueError(f"Could not extract page ID from URL: {url}")
    
    base_url = f"{parsed.scheme}://{parsed.netloc}"
    
    # Fetch the main page
    try:
        page = confluence_get_page(page_id, include_body=True)
    except MCPToolError as e:
        logger.error(f"Failed to fetch page {page_id} via MCP: {e}")
        raise RuntimeError(f"Failed to fetch page {page_id} via MCP: {e}")
    
    documents = []
    
    # Convert page to document format
    body_html = page.get("body_html", page.get("body", ""))
    title = page.get("title", f"Page {page_id}")
    version = page.get("version", 0)
    space_key = page.get("space", "")
    
    # Check if this is Confluence storage format (contains ac:link tags)
    is_storage_format = "<ac:" in body_html or "<ri:" in body_html
    
    if is_storage_format:
        # For storage format, extract page links with their space keys
        import re
        # Extract both space key and title from links
        page_links = re.findall(r'<ri:page ri:space-key="([^"]+)" ri:content-title="([^"]+)"', body_html)
        logger.info(f"Found {len(page_links)} page links in storage format")
        
        # Add the page itself as a document (even with minimal content)
        doc_text = f"# {title}\n\n"
        doc_text += f"Source: Confluence page {page_id} (space: {space_key}, version: {version})\n\n"
        doc_text += f"This page contains links to: {', '.join([f'{t} ({s})' for s, t in page_links]) if page_links else 'No links'}\n"
        doc_text += f"Page content is in Confluence storage format.\n"
        
        documents.append({
            "content": doc_text,
            "metadata": {
                "source": f"{base_url}/pages/{page_id}",
                "title": title,
                "space": space_key,
                "page_id": str(page_id),
                "version": version,
                "type": "confluence_page",
                "parent_page_id": None,
            }
        })
        
        # Fetch linked pages using MCP search tool
        for link_space, link_title in page_links:
            try:
                from agentic_cli.mcp_tool_client import call_mcp_tool, _get_confluence_url
                # Use the space from the link, fall back to page's space
                target_space = link_space if link_space else space_key
                cql = f'space="{target_space}" AND title="{link_title}" AND type=page'
                logger.info(f"Searching for linked page: {link_title} in space {target_space}")
                search_results = call_mcp_tool(_get_confluence_url(), "search_confluence_cql", {"cql": cql, "limit": 1})
                
                # If found, recursively fetch the linked page
                if search_results and search_results != "0":
                    import json
                    try:
                        results = json.loads(search_results) if isinstance(search_results, str) else search_results
                        logger.info(f"Search results type: {type(results)}, keys: {list(results.keys()) if isinstance(results, dict) else 'N/A'}")
                        
                        # Handle different response formats
                        page_id = None
                        if isinstance(results, list) and len(results) > 0:
                            page_id = results[0].get("id")
                        elif isinstance(results, dict):
                            # Check for common dict structures
                            if "results" in results and len(results["results"]) > 0:
                                page_id = results["results"][0].get("id")
                            elif "id" in results:
                                page_id = results.get("id")
                            elif "content" in results and len(results["content"]) > 0:
                                page_id = results["content"][0].get("id")
                        
                        if page_id:
                            logger.info(f"Found linked page ID: {page_id}, fetching...")
                            linked_url = f"{base_url}/pages/{page_id}"
                            linked_docs = _parse_confluence_tree_mcp(
                                linked_url,
                                include_children=True,
                                max_depth=max_depth - 1,
                                include_attachments=include_attachments,
                                current_depth=current_depth + 1,
                            )
                            documents.extend(linked_docs)
                            logger.info(f"Added {len(linked_docs)} documents from linked page {link_title}")
                        else:
                            logger.warning(f"Could not extract page ID from search results for {link_title}")
                    except json.JSONDecodeError as e:
                        logger.warning(f"Failed to parse search results for {link_title}: {e}")
                else:
                    logger.warning(f"Search returned no results for {link_title}")
            except Exception as e:
                logger.error(f"Failed to fetch linked page {link_title}: {e}", exc_info=True)
    else:
        # Standard HTML format
        clean_text = _html_to_text(body_html)
    
    if not is_storage_format and clean_text.strip():
        doc_text = f"# {title}\n\n"
        doc_text += f"Source: Confluence page {page_id} (space: {space_key}, version: {version})\n\n"
        doc_text += clean_text
        
        documents.append({
            "content": doc_text,
            "metadata": {
                "source": f"{base_url}/pages/{page_id}",
                "title": title,
                "space": space_key,
                "page_id": str(page_id),
                "version": version,
                "type": "confluence_page",
                "parent_page_id": parent_page_id,
            }
        })
        logger.info(f"[DEBUG] Page {title} (page_id={page_id}, parent_page_id={parent_page_id})")
    else:
        logger.warning(f"Page \"{title}\" ({page_id}) has no content after HTML cleaning")
    
    # Fetch and parse attachments
    if include_attachments:
        try:
            logger.info(f"Fetching attachments for page \"{title}\" ({page_id})...")
            attachments = confluence_list_attachments(page_id)
            logger.info(f"Attachments result type: {type(attachments)}, length: {len(attachments) if isinstance(attachments, list) else 'N/A'}")
            logger.info(f"Found {len(attachments)} attachment(s) for page \"{title}\" ({page_id})")
            logger.info(f"DEBUG: parent_page_id={parent_page_id} for page {page_id}")
            
            # Extract attachments referenced in page content (current version only)
            referenced_attachments = set()
            if is_storage_format and body_html:
                referenced_attachments = _extract_referenced_attachments(body_html)
                logger.info(f"Found {len(referenced_attachments)} attachment(s) referenced in page content")
            
            # Log attachment filenames for visibility
            if attachments and isinstance(attachments, list):
                attachment_names = [a.get("filename", "unknown") for a in attachments]
                logger.info(f"Attachments: {', '.join(attachment_names[:10])}" + (f" ... and {len(attachment_names) - 10} more" if len(attachment_names) > 10 else ""))
            
            for attachment in attachments:
                filename = attachment.get("filename", "")
                if not filename:
                    logger.debug("Skipping attachment with no filename")
                    continue
                
                # Filter to only process attachments referenced in page content (current version)
                # Skip filter for release pages (detected by title starting with "Release") since they have no content
                is_release_page = title.startswith("Release")
                if not is_release_page and referenced_attachments and filename not in referenced_attachments:
                    logger.debug(f"Skipping attachment not referenced in page content: {filename}")
                    continue
                
                # Check file extension
                ext = Path(filename).suffix.lower()
                if ext not in [".pdf", ".docx", ".doc", ".csv", ".xlsx", ".xls", ".txt", ".json"]:
                    logger.debug(f"Skipping unsupported attachment: {filename} (type: {ext})")
                    continue
                
                logger.info(f"Processing attachment: {filename} (type: {ext})")
                
                # Download attachment using MCP (authenticated)
                attachment_id = attachment.get("id", "")
                if not attachment_id:
                    logger.warning(f"Skipping attachment with no ID: {filename}")
                    continue
                
                try:
                    # Download file via MCP (uses authenticated Confluence client)
                    base64_content = confluence_download_attachment(attachment_id)
                    content = base64.b64decode(base64_content)
                    logger.info(f"Downloaded {filename} ({len(content)} bytes)")
                    
                    # Log if file size seems suspiciously small (might be error page)
                    if len(content) < 1000:
                        logger.warning(f"Downloaded file is very small ({len(content)} bytes), might be an error page")
                        logger.warning(f"First 500 chars: {content[:500]}")
                        continue  # Skip this attachment
                    
                    # Save to temp file
                    with tempfile.NamedTemporaryFile(mode='wb', suffix=ext, delete=False) as f:
                        f.write(content)
                        temp_path = f.name
                    
                    try:
                        # Parse based on file type
                        if ext == ".pdf":
                            attachment_docs = parse_pdf(temp_path)
                        elif ext in [".docx", ".doc"]:
                            attachment_docs = parse_docx(temp_path)
                        elif ext in [".xlsx", ".xls"]:
                            attachment_docs = parse_xlsx(temp_path)
                        elif ext == ".csv":
                            attachment_docs = parse_csv(temp_path)
                        elif ext in [".txt", ".md"]:
                            attachment_docs = parse_text(temp_path)
                        elif ext == ".json":
                            attachment_docs = parse_json(temp_path)
                        else:
                            continue
                        
                        logger.info(f"Parsed {filename} into {len(attachment_docs)} document(s)")
                        
                        # Add attachment metadata and append to documents
                        for doc in attachment_docs:
                            doc["metadata"]["source"] = f"{base_url}/pages/{page_id}/attachments/{filename}"
                            doc["metadata"]["attachment_filename"] = filename
                            doc["metadata"]["attachment_page_id"] = str(page_id)
                            doc["metadata"]["parent_page_id"] = parent_page_id
                            doc["metadata"]["type"] = "confluence_attachment"
                            doc["metadata"]["parent_page_title"] = title
                            # Overwrite title with actual filename instead of temp file name
                            doc["title"] = filename
                            documents.append(doc)
                            logger.info(f"[DEBUG] Attachment {filename} (attachment_page_id={page_id}, parent_page_id={parent_page_id})")
                    finally:
                        # Clean up temp file
                        os.unlink(temp_path)
                        
                except Exception as e:
                    logger.error(f"Failed to process attachment {filename}: {e}")
                    
        except MCPToolError as e:
            logger.error(f"Failed to fetch attachments for page {page_id}: {e}")
    
    # Fetch child pages if include_children is True
    if include_children and max_depth > 0:
        try:
            children = confluence_get_page_children(page_id)
            logger.info(f"Found {len(children)} child page(s) for page {page_id}")
            
            for child in children:
                child_id = child.get("id", "")
                if not child_id:
                    continue
                
                # Recursively fetch child page content and attachments
                child_url = f"{base_url}/pages/{child_id}"
                child_docs = _parse_confluence_tree_mcp(
                    child_url,
                    include_children=True,
                    max_depth=max_depth - 1,
                    include_attachments=include_attachments,
                    current_depth=current_depth + 1,
                    parent_page_id=str(page_id),
                )
                documents.extend(child_docs)
                
        except MCPToolError as e:
            logger.error(f"Failed to fetch child pages for page {page_id}: {e}")
    
    # Debug: log metadata of first few documents
    if documents:
        for i, doc in enumerate(documents[:3]):
            metadata = doc.get("metadata", {})
            print(f"[DEBUG] Document {i} metadata: page_id={metadata.get('page_id')}, parent_page_id={metadata.get('parent_page_id')}, attachment_page_id={metadata.get('attachment_page_id')}, type={metadata.get('type')}")
    
    return documents


def parse_docx(file_path: str) -> List[Dict[str, Any]]:
    """
    Parse DOCX (Word) file and extract text content.
    
    Args:
        file_path: Path to DOCX file
    
    Returns:
        List of document dictionaries
    """
    try:
        from docx import Document
    except ImportError:
        raise ImportError(
            "python-docx not installed. Install with: pip install python-docx"
        )
    
    path = Path(file_path)
    if not path.exists():
        raise FileNotFoundError(f"File not found: {file_path}")
    
    documents = []
    doc = Document(path)
    
    # Extract text from paragraphs
    paragraphs = []
    for para in doc.paragraphs:
        if para.text.strip():
            paragraphs.append(para.text)
    
    # Extract text from tables
    tables_text = []
    for table in doc.tables:
        table_data = []
        for row in table.rows:
            row_data = []
            for cell in row.cells:
                if cell.text.strip():
                    row_data.append(cell.text.strip())
            if row_data:
                table_data.append(" | ".join(row_data))
        if table_data:
            tables_text.append("\n".join(table_data))
    
    # Combine all text
    content = "\n\n".join(paragraphs)
    if tables_text:
        content += "\n\n" + "\n\n".join([f"Table {i+1}:\n{t}" for i, t in enumerate(tables_text)])
    
    if content.strip():
        documents.append({
            "title": path.stem,
            "content": content,
            "metadata": {
                "format": "docx",
                "paragraphs_count": len(paragraphs),
                "tables_count": len(tables_text),
            },
        })
    
    return documents


def parse_xlsx(file_path: str) -> List[Dict[str, Any]]:
    """
    Parse XLSX (Excel) file and extract data from all sheets.
    
    Args:
        file_path: Path to XLSX file
    
    Returns:
        List of document dictionaries (one per sheet)
    """
    try:
        import openpyxl
    except ImportError:
        raise ImportError(
            "openpyxl not installed. Install with: pip install openpyxl"
        )
    
    path = Path(file_path)
    if not path.exists():
        raise FileNotFoundError(f"File not found: {file_path}")
    
    documents = []
    workbook = openpyxl.load_workbook(path, data_only=True)
    
    for sheet_name in workbook.sheetnames:
        sheet = workbook[sheet_name]
        
        # Extract data as CSV-like text
        rows_data = []
        for row in sheet.iter_rows(values_only=True):
            # Filter out completely empty rows
            if any(cell is not None for cell in row):
                row_str = " | ".join(str(cell) if cell is not None else "" for cell in row)
                rows_data.append(row_str)
        
        if rows_data:
            content = f"Sheet: {sheet_name}\n\n" + "\n".join(rows_data)
            
            documents.append({
                "title": f"{path.stem} - {sheet_name}",
                "content": content,
                "metadata": {
                    "format": "xlsx",
                    "sheet": sheet_name,
                    "rows_count": len(rows_data),
                },
            })
    
    return documents


def parse_directory(directory: str, recursive: bool = True) -> List[Dict[str, Any]]:
    """
    Parse all supported files in a directory.
    
    Args:
        directory: Path to directory
        recursive: Whether to search recursively
    
    Returns:
        List of document dictionaries from all files
    """
    path = Path(directory)
    if not path.exists() or not path.is_dir():
        raise ValueError(f"Invalid directory: {directory}")
    
    documents = []
    
    # Define supported extensions and their parsers
    parsers = {
        ".pdf": parse_pdf,
        ".txt": parse_text,
        ".md": parse_text,
        ".csv": parse_csv,
        ".json": parse_json,
        ".docx": parse_docx,
        ".doc": parse_docx,  # Use docx parser for .doc (may not work for old .doc format)
        ".xlsx": parse_xlsx,
        ".xls": parse_xlsx,  # Use xlsx parser for .xls (may not work for old .xls format)
    }
    
    # Find all supported files
    pattern = "**/*" if recursive else "*"
    
    for file_path in path.glob(pattern):
        if file_path.is_file() and file_path.suffix in parsers:
            try:
                parser = parsers[file_path.suffix]
                docs = parser(str(file_path))
                documents.extend(docs)
            except Exception as e:
                # Log error but continue processing other files
                print(f"Error parsing {file_path}: {e}")
    
    return documents


def _extract_code_chunk(content: str, line_start: int, line_end: int) -> str:
    """Extract code chunk between line numbers."""
    if not line_start or not line_end:
        return content
    
    lines = content.split('\n')
    chunk_lines = lines[line_start-1:line_end]
    return '\n'.join(chunk_lines)


def parse_git_repository(
    repo_url: str,
    branch: str = None,
    tag: str = None,
    repo_metadata: Dict[str, Any] = None,
    detailed_analysis: bool = False
) -> List[Dict[str, Any]]:
    """
    Clone and parse Git repository with smart chunking.
    
    Implementation:
    1. Clone repo to temp directory
    2. Use gitingest to generate repository digest (always)
    3. Optionally parse source files with code_analyzer (if detailed_analysis=True)
    4. Extract code structure (functions, classes, imports)
    5. Create documents with persona="developer" metadata
    6. Clean up temp directory
    
    Args:
        repo_url: Git repository URL
        branch: Branch name (optional)
        tag: Tag name (optional)
        repo_metadata: Additional metadata (name, domain, purpose)
        detailed_analysis: If True, analyze individual files for code structure.
                          If False (default), only use gitingest digest (faster).
    
    Returns:
        List of document dictionaries with code structure
    """
    import tempfile
    import shutil
    from pathlib import Path
    
    try:
        import git
    except ImportError:
        raise ImportError(
            "GitPython not installed. Install with: pip install gitpython"
        )
    
    try:
        from gitingest import ingest
    except ImportError:
        raise ImportError(
            "gitingest not installed. Install with: pip install gitingest"
        )
    
    from agentic_cli.kg.code_analyzer import get_analyzer, detect_language
    
    # Create temp directory for cloning
    temp_dir = Path(tempfile.mkdtemp(prefix="agent_cli_git_"))
    
    try:
        print(f"[INFO] Cloning repository: {repo_url}")
        
        # Clone repository
        repo = git.Repo.clone_from(repo_url, temp_dir)
        
        # Checkout specific branch or tag
        if tag:
            print(f"[INFO] Checking out tag: {tag}")
            repo.git.checkout(tag)
        elif branch:
            print(f"[INFO] Checking out branch: {branch}")
            repo.git.checkout(branch)
        
        # Get commit info
        commit_hash = repo.head.commit.hexsha[:8]
        commit_date = repo.head.commit.committed_datetime.isoformat()
        
        # Extract repo name from URL
        repo_name = repo_url.rstrip('/').split('/')[-1].replace('.git', '')
        
        # Prepare metadata
        metadata = repo_metadata or {}
        git_metadata = {
            "source_type": "git",
            "persona": "developer",
            "repo_url": repo_url,
            "repo_name": metadata.get("name", repo_name),
            "domain": metadata.get("domain", ""),
            "purpose": metadata.get("purpose", ""),
            "branch": branch or repo.active_branch.name,
            "tag": tag,
            "commit_hash": commit_hash,
            "commit_date": commit_date,
        }
        
        print(f"[INFO] Repository cloned: {repo_name} ({commit_hash})")
        
        # Generate repository digest using gitingest
        print("[INFO] Generating repository digest with gitingest...")
        try:
            digest = ingest(str(temp_dir))
            # Handle both dict and tuple return types from gitingest
            if isinstance(digest, dict):
                repo_summary = digest.get("summary", "")
            elif isinstance(digest, tuple):
                # gitingest may return (summary, tree) tuple
                repo_summary = digest[0] if digest else ""
            else:
                repo_summary = str(digest) if digest else ""
        except Exception as e:
            print(f"[WARN] gitingest failed: {e}, using fallback")
            repo_summary = f"Repository: {repo_name}"
        
        documents = []
        
        # Add repository overview document
        documents.append({
            "title": f"Repository: {repo_name}",
            "content": repo_summary,
            "metadata": {
                **git_metadata,
                "doc_type": "repository_overview",
                "format": "git"
            }
        })
        
        # Optionally parse source files with detailed code analysis
        if detailed_analysis:
            print("[INFO] Performing detailed code analysis...")
            source_extensions = {".py", ".java", ".sql", ".ddl", ".dml"}
            source_files = []
            
            for ext in source_extensions:
                source_files.extend(temp_dir.rglob(f"*{ext}"))
            
            print(f"[INFO] Found {len(source_files)} source files (Python, Java, SQL/DDL/DML)")
            
            for file_path in source_files:
                try:
                    # Detect language
                    language = detect_language(file_path)
                    if not language:
                        continue
                    
                    # Read file content
                    content = file_path.read_text(encoding='utf-8', errors='ignore')
                    
                    # Get analyzer
                    analyzer = get_analyzer(language)
                    if not analyzer:
                        continue
                    
                    # Analyze code structure
                    analysis = analyzer.analyze_file(file_path, content)
                    
                    # Get relative path from repo root
                    relative_path = file_path.relative_to(temp_dir)
                    
                    # Create document for the file
                    file_metadata = {
                    **git_metadata,
                    "doc_type": "source_file",
                    "language": language,
                    "file_path": str(relative_path),
                    "format": "git"
                    }
                    
                    # Add file overview document
                    file_overview_metadata = {
                    **file_metadata,
                    "imports": analysis.get("imports", []),
                    "classes": [c["name"] for c in analysis.get("classes", [])],
                    "functions": [f["name"] for f in analysis.get("functions", [])],
                    }
                    
                    # For SQL files, add data model metadata
                    if language == "sql":
                        file_overview_metadata.update({
                            "file_type": analysis.get("file_type", "unknown"),
                            "tables": [t["name"] for t in analysis.get("tables", [])],
                            "views": [v["name"] for v in analysis.get("views", [])],
                            "procedures": [p["name"] for p in analysis.get("procedures", [])],
                        })
                    
                    documents.append({
                        "title": f"{repo_name}/{relative_path}",
                        "content": f"{analysis['summary']}\n\n{content[:1000]}...",  # First 1000 chars
                        "metadata": file_overview_metadata
                    })
                    
                    # For SQL files: Create documents for tables, views, procedures
                    if language == "sql":
                        # Create documents for each table
                        for table_info in analysis.get("tables", []):
                            columns_desc = "\n".join([
                                f"  - {col['name']}: {col['type']}" + 
                                (" PRIMARY KEY" if col.get('primary_key') else "") +
                                (" NOT NULL" if not col.get('nullable') else "")
                                for col in table_info.get("columns", [])
                            ])
                            
                            documents.append({
                                "title": f"{repo_name}/{relative_path}::Table::{table_info['name']}",
                                "content": f"Table: {table_info['name']}\n\nColumns ({table_info['column_count']}):\n{columns_desc}",
                                "metadata": {
                                    **file_metadata,
                                    "doc_type": "table",
                                    "table_name": table_info["name"],
                                    "column_count": table_info["column_count"],
                                    "columns": [col["name"] for col in table_info.get("columns", [])],
                                }
                            })
                        documents.append({
                            "title": f"{repo_name}/{relative_path}::Table::{table_info['name']}",
                            "content": f"Table: {table_info['name']}\n\nColumns ({table_info['column_count']}):\n{columns_desc}",
                            "metadata": {
                                **file_metadata,
                                "doc_type": "table",
                                "table_name": table_info["name"],
                                "column_count": table_info["column_count"],
                                "columns": [col["name"] for col in table_info.get("columns", [])],
                            }
                        })
                    
                    # Create documents for each view
                    for view_info in analysis.get("views", []):
                        documents.append({
                            "title": f"{repo_name}/{relative_path}::View::{view_info['name']}",
                            "content": f"View: {view_info['name']}\n\nQuery:\n{view_info.get('query', '')}",
                            "metadata": {
                                **file_metadata,
                                "doc_type": "view",
                                "view_name": view_info["name"],
                            }
                        })
                    
                    # Create documents for each procedure/function
                    for proc_info in analysis.get("procedures", []):
                        documents.append({
                            "title": f"{repo_name}/{relative_path}::{proc_info['type'].title()}::{proc_info['name']}",
                            "content": f"{proc_info['type'].title()}: {proc_info['name']}\n\nParameters: {proc_info.get('parameters', '')}",
                            "metadata": {
                                **file_metadata,
                                "doc_type": proc_info["type"],
                                "procedure_name": proc_info["name"],
                            }
                        })
                    
                    # For Python/Java: Create documents for classes and functions
                    else:
                        # Create documents for each class
                        for class_info in analysis.get("classes", []):
                            class_content = _extract_code_chunk(
                            content,
                            class_info.get("line_start"),
                            class_info.get("line_end")
                            )
                            
                            documents.append({
                            "title": f"{repo_name}/{relative_path}::{class_info['name']}",
                            "content": f"Class: {class_info['name']}\n\n{class_info.get('docstring', '')}\n\n{class_content}",
                            "metadata": {
                                **file_metadata,
                                "doc_type": "class",
                                "class_name": class_info["name"],
                                "methods": [m["name"] for m in class_info.get("methods", [])],
                            }
                            })
                    
                        # Create documents for each function
                        for func_info in analysis.get("functions", []):
                                func_content = _extract_code_chunk(
                                content,
                                func_info.get("line_start"),
                                func_info.get("line_end")
                                )
                            
                                documents.append({
                                "title": f"{repo_name}/{relative_path}::{func_info['name']}",
                                "content": f"Function: {func_info['name']}\n\n{func_info.get('docstring', '')}\n\n{func_content}",
                                "metadata": {
                                    **file_metadata,
                                    "doc_type": "function",
                                    "function_name": func_info["name"],
                                    "args": func_info.get("args", []),
                                }
                            })
                    
                except Exception as e:
                        print(f"[WARN] Error analyzing {file_path}: {e}")
                        continue
        else:
            print("[INFO] Skipping detailed code analysis (using gitingest digest only)")
        
        print(f"[INFO] Parsed {len(documents)} documents from repository")
        
        return documents
        
    finally:
        # Clean up temp directory
        print(f"[INFO] Cleaning up temp directory: {temp_dir}")
        shutil.rmtree(temp_dir, ignore_errors=True)
